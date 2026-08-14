from __future__ import annotations

import re
import uuid
from datetime import datetime
from io import BytesIO
from pathlib import Path

from fastapi import APIRouter, File, HTTPException, UploadFile

from ..config import settings

router = APIRouter(prefix="/api/uploads", tags=["uploads"])

MAX_UPLOAD_BYTES = 10 * 1024 * 1024
UPLOAD_DIR = "uploads"
TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".csv",
    ".json",
    ".py",
    ".js",
    ".jsx",
    ".ts",
    ".tsx",
    ".html",
    ".css",
    ".yaml",
    ".yml",
    ".log",
    ".prd",
}


@router.post("")
async def upload_context_file(file: UploadFile = File(...)) -> dict[str, object]:
    raw_name = Path(file.filename or "uploaded-file").name
    suffix = Path(raw_name).suffix.lower()
    if suffix not in {".pdf", ".docx"} and suffix not in TEXT_EXTENSIONS:
        raise HTTPException(status_code=400, detail="Only PDF, DOCX, and text/code files are supported")

    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="Uploaded file is empty")
    if len(data) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="Uploaded file is larger than 10 MB")

    if suffix == ".pdf":
        text = _extract_pdf_text(data)
    elif suffix == ".docx":
        text = _extract_docx_text(data)
    else:
        text = _decode_text(data)

    text = text.strip()
    if not text:
        raise HTTPException(status_code=400, detail="No readable text could be extracted")

    context_path = _write_context_file(raw_name, text)
    prd_analysis = _prd_analysis(raw_name, text)
    return {
        "name": raw_name,
        "path": context_path,
        "chars": len(text),
        "words": len(re.findall(r"\S+", text)),
        "lines": len(text.splitlines()),
        "bytes": len(data),
        "kind": "pdf" if suffix == ".pdf" else "docx" if suffix == ".docx" else "text",
        "extension": suffix.lstrip(".") or "text",
        "summary": _upload_summary(raw_name, text),
        "preview": _preview_text(text),
        "suggested_prompts": _suggested_prompts(raw_name, suffix, prd_analysis),
        "prd_analysis": prd_analysis,
    }


def _write_context_file(original_name: str, text: str) -> str:
    safe_name = _safe_name(original_name)
    stamp = datetime.utcnow().strftime("%Y%m%d-%H%M%S")
    token = uuid.uuid4().hex[:8]
    rel_path = f"{UPLOAD_DIR}/{stamp}-{token}-{safe_name}.txt"
    target = settings.workdir_path / rel_path
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(
        f"Uploaded source: {original_name}\n\n{text}",
        encoding="utf-8",
    )
    return rel_path


def _safe_name(name: str) -> str:
    stem = Path(name).stem or "uploaded"
    stem = re.sub(r"[^A-Za-z0-9_.-]+", "-", stem).strip("-_.")
    return (stem or "uploaded")[:80]


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-16", "cp1252"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _extract_pdf_text(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise HTTPException(
            status_code=500,
            detail="PDF support needs pypdf. Run: pip install -r requirements.txt",
        ) from exc

    reader = PdfReader(BytesIO(data))
    pages = []
    for index, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages.append(f"[Page {index}]\n{page_text.strip()}")
    return "\n\n".join(pages)



def _extract_docx_text(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise HTTPException(
            status_code=500,
            detail="DOCX support needs python-docx. Run: pip install -r requirements.txt",
        ) from exc

    document = Document(BytesIO(data))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_lines: list[str] = []
    for table_index, table in enumerate(document.tables, start=1):
        table_lines.append(f"[Table {table_index}]")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                table_lines.append(" | ".join(cells))
    return "\n\n".join(paragraphs + table_lines)


def _suggested_prompts(name: str, suffix: str, prd_analysis: dict[str, object] | None = None) -> list[str]:
    label = "this uploaded file"
    if suffix == ".pdf":
        label = "this uploaded PDF"
    elif suffix == ".docx":
        label = "this uploaded Word document"
    if prd_analysis:
        return [
            f"Read {label} and confirm the roles, entities, workflows, database tables, pages, and security rules. Do not create files yet.",
            f"Build Step 1 from {label}: create the requirements analysis and roadmap files only.",
            f"Build Step 2 from {label}: create the file plan and first verified MVP scaffold.",
        ]
    return [
        f"Summarize {label} and list the main requirements.",
        f"Find action items and missing details in {label}.",
        f"Use {label} as context and create a step-by-step build plan.",
    ]

def _upload_summary(name: str, text: str, max_chars: int = 180) -> str:
    first = " ".join(text.strip().split())
    if len(first) > max_chars:
        first = first[:max_chars].rstrip() + "..."
    return f"{name}: {first}"


def _preview_text(text: str, max_chars: int = 500) -> str:
    compact = " ".join(text.strip().split())
    if len(compact) > max_chars:
        return compact[:max_chars].rstrip() + "..."
    return compact

def _prd_analysis(name: str, text: str) -> dict[str, object] | None:
    lower = f"{name}\n{text}".lower()
    if "prd" not in lower and "product requirement" not in lower and "requirements" not in lower:
        return None
    categories = {
        "roles": ["guest", "buyer", "seller", "admin", "user", "manager", "courier", "staff"],
        "entities": ["user", "category", "item", "product", "bid", "order", "cart", "payment", "store", "audit", "student", "customer"],
        "workflows": ["register", "login", "search", "checkout", "auction", "bid", "cod", "cash on delivery", "approval", "moderation", "delivery"],
        "database_tables": ["users", "categories", "items", "products", "bids", "orders", "sessions", "audit", "cart", "stores"],
        "pages": ["homepage", "dashboard", "product page", "admin", "seller", "buyer", "login", "register", "checkout", "orders"],
        "security_rules": ["password", "hash", "rate limit", "otp", "tls", "validation", "fraud", "auth", "role-based"],
    }
    found: dict[str, list[str]] = {}
    for key, terms in categories.items():
        matches = [term for term in terms if term in lower]
        found[key] = matches[:10]
    score = sum(len(values) for values in found.values())
    if score < 3:
        return None
    return {
        "is_prd": True,
        "message": "I found roles, entities, workflows, database tables, pages, and security rules. Do you want me to build Step 1?",
        "found": found,
        "next_step_prompt": f"Build Step 1 from uploaded PRD {name}: create requirements analysis, roadmap, file plan, and verification checklist. Do not build final files yet.",
    }
